#!/usr/bin/env python3
"""
生成 ATS 友好的 Word 简历
用法: python3 gen_docx.py
输出: 何伟杰-简历.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Anthropic 色系
ACCENT = RGBColor(0xD9, 0x77, 0x57)  # 陶土色
TEXT = RGBColor(0x1A, 0x19, 0x18)
TEXT_SUB = RGBColor(0x5E, 0x5D, 0x59)
TEXT_MUTED = RGBColor(0x87, 0x86, 0x7F)
BORDER_COLOR = "E8E6DC"


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_bottom_border(paragraph):
    """给段落加底部边线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), BORDER_COLOR)
    bottom.set(qn('w:space'), '4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run(paragraph, text, bold=False, size=10, color=TEXT, font_name=None):
    """添加一段文字"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def set_paragraph_spacing(paragraph, before=0, after=2, line=1.15):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_section_title(doc, title):
    """添加段落标题（如 SKILLS、EXPERIENCE）"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_bottom_border(p)
    add_run(p, title, bold=True, size=9, color=TEXT_MUTED)
    return p


def add_item_header(doc, title, role, date):
    """添加经历/项目标题行，日期用 tab stop 右对齐"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=1)

    # 设置右对齐 tab stop
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9638')  # ~17cm，A4 页面宽度减去左右边距
    tabs.append(tab)
    pPr.append(tabs)

    add_run(p, title, bold=True, size=11, color=ACCENT)
    add_run(p, f" | {role}", bold=False, size=10, color=TEXT)
    add_run(p, "\t", size=10)  # tab 跳到右边
    add_run(p, date, size=10, color=TEXT_MUTED)
    return p


def add_meta(doc, text):
    """添加灰色描述行"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2)
    add_run(p, text, size=9, color=TEXT_MUTED)
    return p


def add_bullet(doc, segments):
    """
    添加 bullet point
    segments: list of (text, bold) tuples
    """
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before=0, after=1, line=1.2)
    p.paragraph_format.left_indent = Inches(0.25)
    for text, bold in segments:
        add_run(p, text, bold=bold, size=9.5, color=TEXT_SUB if not bold else TEXT)
    return p


def main():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ═══ 姓名 ═══
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2)
    add_run(p, "何伟杰", bold=True, size=22, color=TEXT)

    # ═══ 联系方式 ═══
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4)
    info = "+86 13544248446 | heweijie0306@gmail.com | github.com/heweijie0306 | 深圳"
    add_run(p, info, size=9, color=TEXT_SUB)

    # ═══ Professional Summary ═══
    add_section_title(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4)
    segments = [
        ("具备 ", False), ("AI Agent 全栈工程落地", True), ("与", False),
        ("AI 原生产品架构", True), ("设计能力的研发工程师。就职于 AI 设计 SaaS 创业公司（A-B轮），作为核心骨干直接向 CEO 汇报。具备", False),
        ("生产级 Agent 框架设计与落地", True), ("经验。曾独立发掘业务痛点，主导与参与多项架构级创新，驱动产品注册付费转化率增长超 ", False),
        ("150%", True), ("。", False),
    ]
    for text, bold in segments:
        add_run(p, text, bold=bold, size=10, color=TEXT)

    # ═══ Skills ═══
    add_section_title(doc, "SKILLS")
    skills = [
        ("Agent Infra", "Context Engineering, ReAct Loop, Tool Calling, MCP 协议, RAG pipeline, Agent Sandbox, Skills"),
        ("LLMs", "Prompt Engineering, Transformer 架构, 模型选型与评估, LangChain / LangGraph, OpenAI/Claude SDK"),
        ("Tech Stack", "TypeScript / Python, React / Node.js / FastAPI, HTML/CSS, 向量数据库"),
        ("General", "从 0 到 1 MVP 快速验证, 全英工作环境沟通能力"),
    ]
    for label, tags in skills:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=1)
        add_run(p, f"{label}    ", bold=True, size=9, color=TEXT)
        add_run(p, tags, size=9, color=TEXT_SUB)

    # ═══ Experience ═══
    add_section_title(doc, "EXPERIENCE")
    add_item_header(doc, "深圳灵图创新科技（Pageon）", "AI Agent 算法/应用工程师", "2024/09 – Present")
    add_meta(doc, "A-B轮 AI 设计 SaaS 公司，直接汇报给 CEO，带领并指导 2 名研发实习生。")
    add_bullet(doc, [
        ("早期负责基于 ", False), ("Python/FastAPI", True), (" 的 AI 生成服务后端研发；后续在 ", False),
        ("TypeScript 全栈", True), ("架构下，主要负责核心 AI PPT 生成平台的 ", False),
        ("Agent System", True), ("（自研到 ", False), ("Claude Code / OpenClaw", True),
        (" 体系）与底层数据架构的研发与演进，主导和参与的架构重构推动注册付费转化率从 ", False),
        ("~2% 跃升至 5%+", True), ("。", False),
    ])

    # ═══ Projects ═══
    add_section_title(doc, "PROJECTS")

    # 项目 1
    add_item_header(doc, "Pageon PPT · 数据结构架构重构", "核心发起人 & 方案 Owner", "2025/11 – 2026/01")
    add_meta(doc, "针对旧版 XML + React 方案视觉效果受限的问题，主动发起底层数据结构重构。")
    add_bullet(doc, [
        ("重构数据结构：", True), ("独立发现并提出 ", False), ("HTML 直出 + 前端后处理", True),
        (" 方案——通过 Context Engineering 引导大模型直接输出原生 HTML/CSS 以替代 XML 渲染方案，同时基于 DOM Parser + tiptap 实现富文本编辑与拖拽，兼顾视觉效果与编辑能力。", False),
    ])
    add_bullet(doc, [
        ("推动生产落地：", True), ("独立完成架构验证 Demo 并向团队演示，方案获团队采纳并确立为产品未来技术路线。上线后转化率稳定在 ", False),
        ("5%+", True), ("，视觉效果与 Agent 能力超越竞品 Gamma。同时拓宽产品能力边界，成功支持 Poster、Infographic、Social Media 等多尺寸平面设计场景。", False),
    ])

    # 项目 2
    add_item_header(doc, "Pageon MG视频 · 动画视频生成 Agent", "核心研发", "2026/01 – Present")
    add_meta(doc, "为产品新增 AI 动画视频生成能力，基于 OpenClaw（PiAgent） + Vercel just-bash 搭建完全 Serverless 的 Web 端 Agent 运行时。")
    add_bullet(doc, [
        ("搭建 Agent 基座：", True), ("基于 ", False), ("PiAgent SDK", True),
        (" + ", False), ("just-bash", True),
        (" 构建 Serverless Web 端 Agent 运行时，零容器零 WASM；开发 validate（", False),
        ("Babel", True), (" 编译 + 静态校验）、create_image、TTS 等自定义 Tool，Agent 生成 TSX 代码即可实时预览。", False),
    ])
    add_bullet(doc, [
        ("驱动业务与交付：", True), ("设计 ", False), ("Skills", True),
        (" 系统，将 10 类动画场景（3D、转场、运动设计等）沉淀为领域知识文档，运行时注入 Prompt 驱动 Agent 按场景生成代码；实现多场景编排与 ", False),
        ("AWS Lambda", True), (" 云端渲染导出。（内测阶段）", False),
    ])

    # 项目 3
    add_item_header(doc, "Pageon PPT · Page Agent Infra", "核心研发", "2025/07 – Present")
    add_meta(doc, "为拓宽 AI Chat 的编辑能力与鲁棒性，重构 PPT Agent 基础设施，经历从自研到 Claude Code Agent SDK 的完整演进。")
    add_bullet(doc, [
        ("搭建 Agent 核心工具链：", True), ("自研 ", False), ("Tool Calling", True),
        (" 与 ", False), ("ReAct Loop", True),
        (" 闭环，实现分页读取、文件编辑等核心 Tool，内置 HITL 中断与自动纠错；针对大模型 diff 不稳定问题，设计 ", False),
        ("7 种 Search & Replace 策略", True), ("（含 XML 结构匹配），显著降低编辑失败率。", False),
    ])
    add_bullet(doc, [
        ("推动 Agent 体系演进：", True), ("后期推动 Infra 从自研迁移至 ", False),
        ("Claude Code Agent SDK", True), ("，基于其原生 File System、Bash Tool 重构适配层；同步基于 ", False),
        ("MCP 协议", True), ("开发 stdio Server 并支持 HTTP MCP 接入用户云端数据源（GitHub、Notion 等）。", False),
    ])

    # ═══ Education ═══
    add_section_title(doc, "EDUCATION")
    add_item_header(doc, "宾夕法尼亚州立大学 (Penn State)", "电气工程 · 硕士", "2022/09 – 2024/05")
    add_meta(doc, "GPA: 3.7/4.0 | State College, PA, USA")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=1)
    add_run(p, "研究方向：", bold=True, size=9, color=TEXT_MUTED)
    add_run(p, "图神经网络（GNNs/GATs）、深度学习（PyTorch/TensorFlow）、计算机视觉（CNN）", size=9, color=TEXT_MUTED)

    add_item_header(doc, "华北理工大学", "电子信息工程 · 本科", "2018/09 – 2022/06")
    add_meta(doc, "GPA: 3.5/4.0 | 唐山")

    # 保存
    out = "/Users/iart_weijie/fupan/何伟杰-简历.docx"
    doc.save(out)
    print(f"✅ 已生成: {out}")


if __name__ == "__main__":
    main()
